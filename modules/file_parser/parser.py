import os
from utils.logger import get_logger
from utils.config import SUPPORTED_FILE_TYPES

logger = get_logger()


def parse_txt(file_path: str) -> str:
    try:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        logger.warning(f"TXT文件编码识别失败: {file_path}")
        return ""
    except Exception as e:
        logger.error(f"TXT解析失败: {file_path}, 错误: {e}")
        return ""


def parse_pdf(file_path: str) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        texts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
        text = "\n".join(texts)
        logger.info(f"PDF解析成功: {file_path}, 共{len(reader.pages)}页")
        return text
    except ImportError:
        logger.error("PyPDF2未安装，请执行: pip install PyPDF2")
        return ""
    except Exception as e:
        logger.error(f"PDF解析失败: {file_path}, 错误: {e}")
        return ""


def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        text = "\n".join(texts)
        logger.info(f"DOCX解析成功: {file_path}, 共{len(doc.paragraphs)}段")
        return text
    except ImportError:
        logger.error("python-docx未安装，请执行: pip install python-docx")
        return ""
    except Exception as e:
        logger.error(f"DOCX解析失败: {file_path}, 错误: {e}")
        return ""


def parse_file(file_path: str) -> str:
    if not os.path.exists(file_path):
        logger.error(f"文件不存在: {file_path}")
        return ""

    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext not in SUPPORTED_FILE_TYPES:
        logger.error(f"不支持的文件类型: {ext}, 支持: {SUPPORTED_FILE_TYPES}")
        return ""

    parser_map = {
        "txt": parse_txt,
        "pdf": parse_pdf,
        "docx": parse_docx,
    }

    parser = parser_map.get(ext)
    if parser is None:
        logger.error(f"无对应解析器: {ext}")
        return ""

    text = parser(file_path)
    if not text.strip():
        logger.warning(f"文件解析结果为空: {file_path}")
        return ""

    logger.info(f"文件解析完成: {file_path}, 文本长度: {len(text)}")
    return text


def parse_uploaded_file(uploaded_file) -> str:
    try:
        import tempfile
        suffix = os.path.splitext(uploaded_file.name)[1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getbuffer())
            tmp_path = tmp.name
        text = parse_file(tmp_path)
        os.unlink(tmp_path)
        return text
    except Exception as e:
        logger.error(f"上传文件解析失败: {uploaded_file.name}, 错误: {e}")
        return ""
